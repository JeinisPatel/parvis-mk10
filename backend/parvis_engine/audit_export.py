"""
PARVIS — Audit Report Export Functions
Garamond Word doc + fpdf2 PDF with logo watermark
"""
import io, os, re
from datetime import datetime


def _clean_pdf_text(text: str) -> str:
    """Replace non-Latin-1 characters with ASCII equivalents for fpdf2."""
    replacements = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u00b7": ".",    # middle dot
        "\u2022": "*",    # bullet
        "\u2500": "-",    # box drawing horizontal
        "\u2714": "[OK]", # check mark
        "\u26a0": "(!)",  # warning
        "\u00a9": "(c)",  # copyright
        "\u00d7": "x",    # multiply
        "\u2593": "#",    # block
        "\u00b2": "2",    # superscript 2
        "\u27e9": ">",    # math bracket
        "\u27e8": "<",
        "\u03c8": "psi",  # psi
        "\u03c1": "rho",  # rho
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Strip remaining non-latin1
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _logo_paths():
    """Return available logo file paths in order of preference."""
    candidates = {
        "light": [
            "ethical_ai_logo_light.png",
            "parvis/ethical_ai_logo_light.png",
            "/mount/src/parvis/ethical_ai_logo_light.png",
        ],
        "pdf": [
            "ethical_ai_logo_pdf.png",
            "parvis/ethical_ai_logo_pdf.png",
            "/mount/src/parvis/ethical_ai_logo_pdf.png",
        ],
        "orig": [
            "ethical_ai_logo.png",
            "parvis/ethical_ai_logo.png",
            "/mount/src/parvis/ethical_ai_logo.png",
        ],
    }
    result = {}
    for key, paths in candidates.items():
        for p in paths:
            if os.path.exists(p):
                result[key] = p
                break
    return result


def build_docx(Pa, da, bla, cG, cS, mx, cr_rec, cr_doc_adj, doc_adj,
               scefw, conn, qdiags, NODE_META) -> bytes:
    """
    Professional Word document audit report.
    Font: Garamond throughout.
    Title: P. [logo] .R.V.I.S inline, matching Streamlit header.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    logos = _logo_paths()
    FONT = "Garamond"

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

    # ── Header: table layout — text left, logo right, navy rule ─────────────
    header = section.header
    # Remove default empty paragraph
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    htbl = header.add_table(1, 2, width=Cm(15.4))
    htbl.style = "Table Grid"
    tblPr = htbl._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Left cell: document title text
    lc = htbl.rows[0].cells[0]
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run("P.A.R.V.I.S  ·  Audit Report")
    lr.font.name = FONT; lr.font.size = Pt(8.5)
    lr.font.color.rgb = RGBColor.from_string("888888")

    # Right cell: logo, vertically matching text
    rc = htbl.rows[0].cells[1]
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logos.get("light"):
        rr = rp.add_run()
        rr.add_picture(logos["light"], height=Pt(18))

    # Bottom rule paragraph
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(1)
    rule_p.paragraph_format.space_after  = Pt(0)
    pPr = rule_p.paragraph_format._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), "1B2A4A")
    pBdr.append(bot); pPr.append(pBdr)

    # ── Footer: page numbers ───────────────────────────────────────────────────
    footer = section.footer
    ftp = footer.paragraphs[0]
    ftp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr = ftp.add_run("P.A.R.V.I.S Audit Report  |  Page ")
    ftr.font.size = Pt(8); ftr.font.name = FONT
    ftr.font.color.rgb = RGBColor.from_string("888888")
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run_pn = OxmlElement("w:r")
    run_pn.append(fldChar1); run_pn.append(instrText); run_pn.append(fldChar2)
    ftp._p.append(run_pn)

    # ── Helper functions ────────────────────────────────────────────────────────
    def add_rule(color="1B2A4A", thickness="8"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr2 = p.paragraph_format._element.get_or_add_pPr()
        pBdr2 = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), thickness)
        b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
        pBdr2.append(b); pPr2.append(pBdr2)
        return p

    def add_section(text, color="1B2A4A"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(11.5)
        r.font.name = FONT
        r.font.color.rgb = RGBColor.from_string(color)
        # Section underline
        pPr3 = p.paragraph_format._element.get_or_add_pPr()
        pBdr3 = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "2"); b.set(qn("w:color"), color)
        pBdr3.append(b); pPr3.append(pBdr3)
        return p

    def add_body(text, color="333333", indent=False, bold=False, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        r.bold = bold
        return p

    def add_field(label, value, lc="185FA5", vc="222222"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.6)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = FONT
        r1.font.color.rgb = RGBColor.from_string(lc)
        r2 = p.add_run(str(value))
        r2.font.size = Pt(10.5); r2.font.name = FONT
        r2.font.color.rgb = RGBColor.from_string(vc)
        return p

    # ── TITLE: P. [logo] .R.V.I.S ─────────────────────────────────────────────
    doc.add_paragraph()  # top spacing
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(4)
    # Fixed line spacing prevents logo from pushing line height unpredictably
    tp.paragraph_format.line_spacing = Pt(44)

    r_p = tp.add_run("P.")
    r_p.bold = True; r_p.font.size = Pt(36)
    r_p.font.name = FONT
    r_p.font.color.rgb = RGBColor.from_string("1B2A4A")

    # Inline logo — sized to match cap height, dist set for vertical centering
    if logos.get("light"):
        r_logo = tp.add_run()
        r_logo.add_picture(logos["light"], width=Pt(34))
        # Adjust vertical position via inline image XML
        from docx.shared import Emu
        drawing = tp._p.findall(".//" + qn("w:drawing"))[-1]
        inline = drawing.find(qn("wp:inline"))
        if inline is not None:
            inline.set("distT", str(int(Pt(2).emu)))
            inline.set("distB", "0")
            inline.set("distL", str(int(Pt(1).emu)))
            inline.set("distR", str(int(Pt(1).emu)))
    else:
        r_dot = tp.add_run("A")
        r_dot.bold = True; r_dot.font.size = Pt(36)
        r_dot.font.name = FONT
        r_dot.font.color.rgb = RGBColor.from_string("1B2A4A")

    r_rest = tp.add_run(".R.V.I.S")
    r_rest.bold = True; r_rest.font.size = Pt(36)
    r_rest.font.name = FONT
    r_rest.font.color.rgb = RGBColor.from_string("1B2A4A")

    # Subtitle: italic system name
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp.paragraph_format.space_after = Pt(2)
    sr = sp.add_run("Probabilistic and Analytical Reasoning Virtual Intelligence System")
    sr.italic = True; sr.font.size = Pt(10.5); sr.font.name = FONT
    sr.font.color.rgb = RGBColor.from_string("555555")

    # Authorship line
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ip.paragraph_format.space_after = Pt(10)
    for part, col in [
        ("Jeinis Patel, PhD Candidate and Barrister", "A32D2D"),
        ("  \u00b7  University of London", "185FA5"),
        ("  \u00b7  Ethical AI Initiative", "1B5E20"),
        (f"  \u00b7  {datetime.now().strftime('%d %B %Y  %H:%M')}", "777777"),
        ("  \u00b7  \u00a9 2026 Jeinis Patel", "999999"),
    ]:
        ir = ip.add_run(part)
        ir.font.size = Pt(9.5); ir.font.name = FONT
        ir.font.color.rgb = RGBColor.from_string(col)

    add_rule("1B2A4A", "10")

    # ── Inference output ─────────────────────────────────────────────────────
    add_section("INFERENCE OUTPUT", "1B2A4A")
    risk_col = "A32D2D" if da >= 0.55 else "BA7517" if da >= 0.35 else "3B6D11"
    rp = doc.add_paragraph()
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after  = Pt(4)
    rr1 = rp.add_run("Node 20  \u2014  DO Designation Risk:   ")
    rr1.bold = True; rr1.font.size = Pt(14); rr1.font.name = FONT
    rr1.font.color.rgb = RGBColor.from_string("1B2A4A")
    rr2 = rp.add_run(f"{da*100:.2f}%   [{bla.upper()}]")
    rr2.bold = True; rr2.font.size = Pt(14); rr2.font.name = FONT
    rr2.font.color.rgb = RGBColor.from_string(risk_col)
    add_body(
        "This figure represents the posterior probability of Dangerous Offender designation "
        "given all upstream evidence, corrections, and doctrinal adjustments applied. "
        "It models DESIGNATION RISK \u2014 not intrinsic dangerousness. This distinction is "
        "the thesis's central normative contribution.",
        color="444444")

    # ── Doctrinal framework ──────────────────────────────────────────────────
    add_section("DOCTRINAL FRAMEWORK  \u2014  THE TETRAD", "185FA5")
    for cite in [
        "R v Gladue [1999] 1 SCR 688",
        "R v Ipeelee [2012] SCC 13",
        f"R v Morris 2021 ONCA 680 (para 97)  \u2014  Framework: {scefw.upper()}  \u00b7  Connection: {conn.upper()}  \u00b7  Multiplier: {mx:.2f}",
        "R v Ellis 2022 BCCA 278",
        "Ewert v Canada [2018] SCC 30",
        "R v Boutilier 2017 SCC 64",
        "R v Natomagan 2022 ABCA 48",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.6)
        r = p.add_run(f"\u25b6  {cite}")
        r.italic = True; r.font.size = Pt(10.5); r.font.name = FONT
        r.font.color.rgb = RGBColor.from_string("185FA5")

    # ── Gladue factors ────────────────────────────────────────────────────────
    add_section("GLADUE FACTORS", "1B5E20")
    if cG:
        for f in cG:
            add_field(f"[+] {f['l']}", f"Node {f['n']}  (+{f['w']*100:.0f}%)", "1B5E20")
    else:
        add_body("No Gladue factors selected.", color="888888")

    # ── Morris / Ellis SCE ────────────────────────────────────────────────────
    add_section("MORRIS / ELLIS SOCIAL CONTEXT EVIDENCE", "185FA5")
    if cS:
        for f in cS:
            add_field(f"[+] {f['l']}",
                f"Node {f['n']}  (+{f['w']*mx*100:.1f}% after connection weight {mx:.2f})",
                "185FA5")
    else:
        add_body("No Morris/Ellis SCE factors selected.", color="888888")

    # ── Criminal record ───────────────────────────────────────────────────────
    if cr_rec:
        add_section("CALIBRATED CRIMINAL RECORD", "A32D2D")
        esc_i = cr_doc_adj.get("escalation", {})
        add_field("Pattern (Boutilier)", esc_i.get("pattern","--").title(), "A32D2D")
        add_body(esc_i.get("note", ""), color="555555")

        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        tbl.autofit = False
        # Column widths in DXA (1440 = 1 inch)
        widths_cm = [1.8, 6.0, 4.5, 2.2, 1.8]
        for idx, (cell, w) in enumerate(zip(tbl.rows[0].cells, widths_cm)):
            cell.width = Cm(w)
        hdr_labels = ["Year", "Offence", "Seriousness", "Cal. Wt", "Gang"]
        for ci, (cell, lbl) in enumerate(zip(tbl.rows[0].cells, hdr_labels)):
            cell.text = lbl
            r = cell.paragraphs[0].runs[0]
            r.bold = True; r.font.size = Pt(9); r.font.name = FONT
            # Dark header shading
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1B2A4A")
            tcPr.append(shd)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for e in cr_rec:
            row = tbl.add_row()
            vals = [
                str(e["year"]), e["offence"][:45],
                e.get("seriousness_label", "--"),
                f"{e['cal_weight']*100:.0f}%",
                "Yes" if e.get("gang") else "No",
            ]
            for ci, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.text = val
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.name = FONT
        doc.add_paragraph()

    # ── Risk posteriors ────────────────────────────────────────────────────────
    add_section("RISK FACTOR POSTERIORS (Variable Elimination)", "A32D2D")
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    for ci, ht in enumerate(["Node", "Factor", "P(High) %"]):
        tbl2.rows[0].cells[ci].text = ht
        if tbl2.rows[0].cells[ci].paragraphs[0].runs:
            r = tbl2.rows[0].cells[ci].paragraphs[0].runs[0]
            r.bold = True; r.font.size = Pt(9); r.font.name = FONT
    for nid in NODE_META:
        if NODE_META[nid]["type"] == "risk" and nid != 20:
            pval = Pa.get(nid, 0.5) * 100
            row = tbl2.add_row()
            for ci, val in enumerate([f"N{nid}", NODE_META[nid]["name"], f"{pval:.1f}%"]):
                row.cells[ci].text = val
                if row.cells[ci].paragraphs[0].runs:
                    row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
                    row.cells[ci].paragraphs[0].runs[0].font.name = FONT
    doc.add_paragraph()

    # ── Distortion corrections ─────────────────────────────────────────────────
    add_section("SYSTEMIC DISTORTION CORRECTIONS", "185FA5")
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    for ci, ht in enumerate(["Node", "Factor", "P(High) %"]):
        tbl3.rows[0].cells[ci].text = ht
        if tbl3.rows[0].cells[ci].paragraphs[0].runs:
            r = tbl3.rows[0].cells[ci].paragraphs[0].runs[0]
            r.bold = True; r.font.size = Pt(9); r.font.name = FONT
    for nid in NODE_META:
        if NODE_META[nid]["type"] not in ("risk", "output") and nid != 20:
            pval = Pa.get(nid, 0.5) * 100
            row = tbl3.add_row()
            for ci, val in enumerate([f"N{nid}", NODE_META[nid]["name"], f"{pval:.1f}%"]):
                row.cells[ci].text = val
                if row.cells[ci].paragraphs[0].runs:
                    row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
                    row.cells[ci].paragraphs[0].runs[0].font.name = FONT
    doc.add_paragraph()

    # ── QBism diagnostics ─────────────────────────────────────────────────────
    if qdiags:
        add_section("QBISM DIAGNOSTIC REPORT (APPENDIX Q)", "534AB7")
        qd = qdiags
        add_field("Overall flag", qd.get("overall_flag", "--").upper(), "534AB7")
        add_field("Superposition index", f"{qd.get('superposition_index','--')} / 1.0", "534AB7")
        add_body(qd.get("summary", ""), color="444444")
        for check, label in [
            ("prior_contamination",     "Prior Contamination"),
            ("order_effects",           "Order Effects (Non-Commutativity)"),
            ("contextual_interference", "Contextual Interference"),
            ("belief_stasis",           "Belief Stasis / Scalar Collapse"),
        ]:
            d = qd.get(check, {})
            sev = d.get("severity", "none").upper()
            sv_col = "A32D2D" if sev=="HIGH" else "BA7517" if sev=="MODERATE" else "3B6D11"
            flagged = "[FLAGGED]" if d.get("flagged") else "[Clear]"
            add_field(label, f"{sev}  {flagged}", sv_col)

    # ── Footer rule + disclaimer ───────────────────────────────────────────────
    add_rule("A32D2D", "8")
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(6)
    for part, col in [
        ("PARVIS  \u00b7  Research use only  \u00b7  NOT for deployment in live proceedings\n", "A32D2D"),
        ("\u00a9 Jeinis Patel, PhD Candidate and Barrister  \u00b7  Ethical AI Initiative  \u00b7  University of London", "666666"),
    ]:
        fr = fp.add_run(part)
        fr.font.size = Pt(8.5); fr.font.name = FONT; fr.italic = True
        fr.font.color.rgb = RGBColor.from_string(col)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(Pa, da, bla, cG, cS, mx, cr_rec, cr_doc_adj,
              scefw, conn, qdiags, NODE_META) -> bytes:
    """Clean PDF using fpdf2. All special characters sanitised for Latin-1."""
    from fpdf import FPDF
    import os

    logos = _logo_paths()
    pdf_logo = logos.get("pdf") or logos.get("orig")

    class PARVISReport(FPDF):
        def header(self):
            # Navy top bar
            self.set_fill_color(27, 42, 74)
            self.rect(0, 0, 210, 2, "F")
            # Logo right
            if pdf_logo:
                self.image(pdf_logo, x=185, y=3, w=12)
            # Header text
            self.set_xy(14, 3)
            self.set_font("Helvetica", "B", 8)
            self.set_text_color(27, 42, 74)
            self.cell(0, 7, "P.A.R.V.I.S  |  Audit Report", ln=True)
            self.ln(4)

        def footer(self):
            self.set_y(-14)
            self.set_draw_color(27, 42, 74)
            self.set_line_width(0.4)
            self.line(14, self.get_y(), 196, self.get_y())
            self.ln(1)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(140, 140, 140)
            self.cell(0, 5,
                f"PARVIS | Research use only | NOT for deployment in live proceedings | "
                f"(c) Jeinis Patel | Page {self.page_no()}",
                align="C")

    pdf = PARVISReport(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(16, 24, 16)

    # ── Watermark ────────────────────────────────────────────────────────────
    if pdf_logo:
        with pdf.local_context(fill_opacity=0.05, stroke_opacity=0.05):
            pdf.image(pdf_logo, x=55, y=75, w=100)

    def section_head(text, r=27, g=42, b=74):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(r, g, b)
        pdf.cell(0, 7, _clean_pdf_text(text), ln=True)
        pdf.set_draw_color(r, g, b)
        pdf.set_line_width(0.3)
        pdf.line(16, pdf.get_y(), 194, pdf.get_y())
        pdf.ln(3)

    def body(text, r=50, g=50, b=50, indent=True):
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(r, g, b)
        if indent:
            pdf.set_x(22)
        pdf.multi_cell(0, 5.5, _clean_pdf_text(text), ln=True)

    def field(label, value, lr=24, lg=95, lb=165):
        pdf.set_x(22)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(lr, lg, lb)
        label_w = 65
        pdf.cell(label_w, 5.5, _clean_pdf_text(label) + ":", ln=False)
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(0, 5.5, _clean_pdf_text(str(value)), ln=True)

    # ── Title ────────────────────────────────────────────────────────────────
    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 30)
    pdf.set_text_color(27, 42, 74)
    # P. + logo + .R.V.I.S on one visual line
    pdf.cell(20, 14, "P.", ln=False)
    if pdf_logo:
        y_before = pdf.get_y()
        pdf.image(pdf_logo, x=pdf.get_x(), y=y_before+1, w=12)
        pdf.set_xy(pdf.get_x()+14, y_before)
    else:
        pdf.cell(14, 14, "A", ln=False)
    pdf.cell(0, 14, ".R.V.I.S", ln=True)

    pdf.set_font("Helvetica", "I", 10.5)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 6, "Probabilistic and Analytical Reasoning Virtual Intelligence System", ln=True)
    pdf.ln(1)

    # Authorship
    pdf.set_font("Helvetica", "", 9)
    for part, (r, g, b) in [
        ("Jeinis Patel, PhD Candidate and Barrister", (163, 45, 45)),
        ("   |   University of London", (24, 95, 165)),
        ("   |   Ethical AI Initiative", (27, 94, 32)),
        (f"   |   {datetime.now().strftime('%d %B %Y')}", (100, 100, 100)),
    ]:
        pdf.set_text_color(r, g, b)
        pdf.cell(pdf.get_string_width(part)+2, 6, _clean_pdf_text(part), ln=False)
    pdf.ln(2)

    pdf.set_draw_color(27, 42, 74)
    pdf.set_line_width(0.5)
    pdf.line(16, pdf.get_y(), 194, pdf.get_y())
    pdf.ln(5)

    # ── Inference output ─────────────────────────────────────────────────────
    section_head("INFERENCE OUTPUT")
    risk_rgb = (163,45,45) if da>=0.55 else (186,117,23) if da>=0.35 else (59,109,17)
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(*risk_rgb)
    pdf.cell(0, 9, f"Node 20 - DO Designation Risk:  {da*100:.2f}%  [{bla.upper()}]", ln=True)
    body("This figure models DESIGNATION RISK -- not intrinsic dangerousness. "
         "Distortion node corrections reduce the evidentiary weight of upstream risk signals.")
    pdf.ln(2)

    # ── Tetrad ───────────────────────────────────────────────────────────────
    section_head("DOCTRINAL FRAMEWORK - THE TETRAD", 24, 95, 165)
    for cite in [
        "R v Gladue [1999] 1 SCR 688",
        "R v Ipeelee [2012] SCC 13",
        f"R v Morris 2021 ONCA 680 -- {scefw.upper()} | Connection: {conn} | x{mx:.2f}",
        "R v Ellis 2022 BCCA 278",
        "Ewert v Canada [2018] SCC 30",
        "R v Boutilier 2017 SCC 64",
        "R v Natomagan 2022 ABCA 48",
    ]:
        pdf.set_x(22); pdf.set_font("Helvetica","I",9.5)
        pdf.set_text_color(24,95,165)
        pdf.cell(5,5.5,">"); pdf.multi_cell(0,5.5,_clean_pdf_text(cite),ln=True)

    # ── Gladue ───────────────────────────────────────────────────────────────
    section_head("GLADUE FACTORS", 27, 94, 32)
    if cG:
        for f in cG:
            field(f"[+] {f['l']}", f"Node {f['n']}  (+{f['w']*100:.0f}%)", 27,94,32)
    else:
        body("No Gladue factors selected.", 130,130,130)

    # ── Morris/Ellis ──────────────────────────────────────────────────────────
    section_head("MORRIS / ELLIS SOCIAL CONTEXT EVIDENCE", 24, 95, 165)
    if cS:
        for f in cS:
            field(f['l'], f"Node {f['n']}  (+{f['w']*mx*100:.1f}%)", 24,95,165)
    else:
        body("No Morris/Ellis SCE factors selected.", 130,130,130)

    # ── Criminal record ───────────────────────────────────────────────────────
    if cr_rec:
        section_head("CALIBRATED CRIMINAL RECORD", 163, 45, 45)
        esc_i = cr_doc_adj.get("escalation", {})
        field("Pattern (Boutilier)", esc_i.get("pattern","--").title(), 163,45,45)
        body(_clean_pdf_text(esc_i.get("note","")))

        # Table
        col_ws = [18, 72, 42, 22, 20]
        pdf.set_font("Helvetica","B",8.5)
        pdf.set_text_color(255,255,255)
        pdf.set_fill_color(27,42,74)
        pdf.set_x(16)
        for txt, w in zip(["Year","Offence","Seriousness","Weight","Gang"], col_ws):
            pdf.cell(w, 6.5, txt, border=0, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica","",8.5)
        for ei, e in enumerate(cr_rec):
            pdf.set_fill_color(248,250,253) if ei%2==0 else pdf.set_fill_color(255,255,255)
            pdf.set_text_color(40,40,40)
            pdf.set_x(16)
            for txt, w in zip([
                str(e["year"]),
                _clean_pdf_text(e["offence"][:38]),
                _clean_pdf_text(e.get("seriousness_label","--")[:20]),
                f"{e['cal_weight']*100:.0f}%",
                "Yes" if e.get("gang") else "No",
            ], col_ws):
                pdf.cell(w, 6, txt, border=0, fill=True)
            pdf.ln()
        pdf.ln(4)

    # ── Risk posteriors ────────────────────────────────────────────────────────
    section_head("RISK FACTOR POSTERIORS", 163, 45, 45)
    for nid in NODE_META:
        if NODE_META[nid]["type"]=="risk" and nid!=20:
            pval = Pa.get(nid,0.5)*100
            col = (163,45,45) if pval>=60 else (186,117,23) if pval>=35 else (59,109,17)
            field(f"N{nid}  {NODE_META[nid]['name']}", f"{pval:.1f}%", *col)
    pdf.ln(2)

    # ── Distortions ────────────────────────────────────────────────────────────
    section_head("SYSTEMIC DISTORTION CORRECTIONS", 24, 95, 165)
    for nid in NODE_META:
        if NODE_META[nid]["type"] not in ("risk","output") and nid!=20:
            pval = Pa.get(nid,0.5)*100
            field(f"N{nid}  {NODE_META[nid]['name']}", f"{pval:.1f}%", 24,95,165)
    pdf.ln(2)

    # ── QBism ──────────────────────────────────────────────────────────────────
    if qdiags:
        section_head("QBISM DIAGNOSTIC REPORT (APPENDIX Q)", 83, 74, 183)
        qd = qdiags
        field("Overall flag", qd.get("overall_flag","--").upper(), 83,74,183)
        field("Superposition index", f"{qd.get('superposition_index','--')} / 1.0", 83,74,183)
        body(_clean_pdf_text(qd.get("summary","")))

    # ── Footer rule ────────────────────────────────────────────────────────────
    pdf.ln(4)
    pdf.set_draw_color(163,45,45)
    pdf.set_line_width(0.4)
    pdf.line(16, pdf.get_y(), 194, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica","I",8)
    pdf.set_text_color(120,120,120)
    pdf.cell(0,5,
        f"Research use only. Not for deployment in live proceedings. "
        f"Generated {datetime.now().strftime('%d %B %Y %H:%M')} | pgmpy Variable Elimination",
        align="C", ln=True)

    return bytes(pdf.output())
