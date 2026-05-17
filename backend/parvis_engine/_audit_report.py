"""
PARVIS Mk 9 - Audit report builder.

Composes a structured DOCX or PDF audit report from current case state,
suitable for submission as a doctrinally-literate companion document to a
sentencing brief or expert report.

The module is intentionally Mk 9-native: it does NOT depend on the legacy
Mk 8 audit_export module, which uses fpdf2 (Latin-1 only) and opaque
parameter names. Mk 9 uses python-docx + reportlab, both Unicode-native.

Public API:
    build_audit_docx(payload: dict, sections: list[str]) -> bytes
    build_audit_pdf(payload: dict, sections: list[str]) -> bytes

Citation markers in narratives are rendered as bold italic inline text:
    [gladue:66]  ->  Gladue para 66 (bold italic)
    [ipeelee:73] ->  Ipeelee para 73
    [ellis:N]    ->  Ellis (no paragraph)
"""

from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
)


# ============================================================================
# CONSTANTS
# ============================================================================

VERSION = "PARVIS Mk 9.0.0-beta"

CITATION_LABELS = {
    "gladue": "Gladue",
    "ipeelee": "Ipeelee",
    "sharma": "Sharma",
    "morris": "Morris",
    "anderson": "Anderson",
    "ellis": "Ellis",
}

CITATION_PATTERN = re.compile(
    r"\[(gladue|ipeelee|sharma|morris|anderson|ellis):([^\]]+)\]"
)

INDIGENOUS_LABELS = {
    "none": "Not Indigenous-identified",
    "first_nations": "First Nations",
    "metis": "Metis",
    "inuit": "Inuit",
    "other_indigenous": "Indigenous (other)",
}

FULL_CITATIONS = {
    "Gladue": "R v Gladue, [1999] 1 SCR 688",
    "Ipeelee": "R v Ipeelee, 2012 SCC 13",
    "Sharma": "R v Sharma, 2022 SCC 39",
    "Morris": "R v Morris, 2021 ONCA 680",
    "Anderson": "R v Anderson, 2021 NSCA 62",
    "Ellis": "R v Ellis (anti-Black racism sentencing framework, Ontario)",
}

ALL_SECTIONS = [
    "title",
    "executive_summary",
    "profile",
    "documents",
    "intake",
    "risk",
    "gladue",
    "sce",
    "authorities",
]


# ============================================================================
# CITATION MARKER PARSING
# ============================================================================

def parse_citations(text: str):
    """Split text into segments. Each segment is a tuple:
        ("text", content, None)     - plain text
        ("cite", case_label, para)  - citation, para may be None
    """
    if not text:
        return []
    segments = []
    last_end = 0
    for m in CITATION_PATTERN.finditer(text):
        if m.start() > last_end:
            segments.append(("text", text[last_end:m.start()], None))
        marker, para = m.group(1), m.group(2)
        label = CITATION_LABELS.get(marker, marker.title())
        if para.strip().upper() in {"N", ""}:
            segments.append(("cite", label, None))
        else:
            segments.append(("cite", label, para.strip()))
        last_end = m.end()
    if last_end < len(text):
        segments.append(("text", text[last_end:], None))
    return segments


def extract_authorities(*narratives):
    """Pull a deduplicated, sorted list of cited authorities from narratives."""
    found = set()
    for n in narratives:
        if not n:
            continue
        for m in CITATION_PATTERN.finditer(n):
            label = CITATION_LABELS.get(m.group(1), m.group(1).title())
            found.add(label)
    return sorted(found)


# ============================================================================
# DOCX BUILDER
# ============================================================================

def _docx_set_default_font(doc, font_name="Garamond", size=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)


def _docx_add_cited_paragraph(doc, text, justify=True):
    """Add a paragraph containing [gladue:N] etc markers, rendered bold italic."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for kind, value, para in parse_citations(text):
        if kind == "text":
            p.add_run(value)
        else:
            run = p.add_run(f"{value} para {para}" if para else value)
            run.italic = True
            run.bold = True


def _docx_add_section_title(doc, title):
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.name = "Garamond"
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def _docx_add_kv_row(table, key, value):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[1].text = value or "-"


def _docx_section_title_page(doc, payload):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PARVIS")
    run.font.name = "Garamond"
    run.font.size = Pt(48)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Probabilistic and Analytical Reasoning Virtual Intelligence System")
    run.font.name = "Garamond"
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    case_ref = payload.get("case_reference") or "Case reference not set"
    profile = payload.get("profile", {}) or {}
    given = profile.get("givenName", "") or ""
    family = profile.get("familyName", "") or ""
    full_name = f"{given} {family}".strip() or "Subject"
    jurisdiction = profile.get("jurisdiction", "") or ""
    court = profile.get("court", "") or ""
    charge = profile.get("primaryCharge", "") or ""
    generated = payload.get("generated_at") or datetime.utcnow().isoformat()

    audit_p = doc.add_paragraph()
    audit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = audit_p.add_run("Audit Report")
    run.font.name = "Garamond"
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph()

    forum_line = f"{court} - {jurisdiction}".strip(" -")
    for line, big in [(case_ref, True), (full_name, True), (charge, False), (forum_line, False)]:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Garamond"
        run.font.size = Pt(14 if big else 12)

    for _ in range(6):
        doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run(f"Generated {generated[:10]} | {VERSION}")
    run.font.name = "Garamond"
    run.font.size = Pt(10)
    run.italic = True

    doc.add_page_break()


def _docx_section_executive_summary(doc, payload):
    _docx_add_section_title(doc, "Executive summary")

    inference = payload.get("inference", {}) or {}
    posterior = inference.get("do_posterior")

    if posterior is None:
        _docx_add_cited_paragraph(
            doc,
            "Posterior probability of designation has not been computed for this case. "
            "Please refer to the Risk State section for current evidence values."
        )
    else:
        pct = posterior * 100
        zone = (
            "above the beyond-reasonable-doubt threshold (>=87%)"
            if pct >= 87
            else "below the beyond-reasonable-doubt threshold (<87%)"
        )
        text = (
            f"Posterior probability that the predicate findings for designation "
            f"are established on the evidence currently configured: {pct:.1f}%. "
            f"This is {zone}. The standard of proof at a s.753 hearing is beyond "
            f"a reasonable doubt; PARVIS operationalises BRD at 87% on the basis "
            f"of empirical literature on judicial interpretation of that standard "
            f"(R v Lifchus and R v Starr describe BRD qualitatively, without "
            f"specifying a probability). PARVIS does not predict; it surfaces "
            f"the structure of the inference. This report is intended as a "
            f"doctrinally literate companion to the substantive submissions of counsel."
        )
        _docx_add_cited_paragraph(doc, text)

    drivers_up = inference.get("drivers_up", []) or []
    drivers_down = inference.get("drivers_down", []) or []

    if drivers_up or drivers_down:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Drivers")
        run.bold = True

        if drivers_up:
            p = doc.add_paragraph()
            run = p.add_run("Increasing posterior: ")
            run.italic = True
            p.add_run(", ".join(
                f"{d.get('label', d.get('nodeId', '?'))} ({d.get('delta', 0):+.1%})"
                for d in drivers_up[:5]
            ))
        if drivers_down:
            p = doc.add_paragraph()
            run = p.add_run("Decreasing posterior: ")
            run.italic = True
            p.add_run(", ".join(
                f"{d.get('label', d.get('nodeId', '?'))} ({d.get('delta', 0):+.1%})"
                for d in drivers_down[:5]
            ))

    doc.add_page_break()


def _docx_section_profile(doc, payload):
    _docx_add_section_title(doc, "Case profile")
    profile = payload.get("profile", {}) or {}

    if not profile:
        _docx_add_cited_paragraph(doc, "No case profile data captured.")
        doc.add_page_break()
        return

    table = doc.add_table(rows=0, cols=2)
    table.autofit = True

    fields = [
        ("Case reference", profile.get("caseReference")),
        ("Given name(s)", profile.get("givenName")),
        ("Family name", profile.get("familyName")),
        ("Date of birth", profile.get("dateOfBirth")),
        ("Pronouns", profile.get("pronouns")),
        ("Indigenous identity", INDIGENOUS_LABELS.get(profile.get("indigenousIdentity", "none"), "-")),
        ("Nation / community", profile.get("nationCommunity")),
        ("Place of origin", profile.get("placeOfOrigin")),
        ("Current residence", profile.get("currentResidence")),
        ("FASD diagnosis", profile.get("fasdDiagnosis")),
        ("Cognitive assessment", profile.get("cognitiveAssess")),
        ("Mental health diagnosis", profile.get("mentalHealthDx")),
        ("Primary charge", profile.get("primaryCharge")),
        ("Additional charges", profile.get("additionalCharges")),
        ("s.753 application", profile.get("s753Application")),
        ("Crown position", profile.get("crownPosition")),
        ("Defence position", profile.get("defencePosition")),
        ("Jurisdiction", profile.get("jurisdiction")),
        ("Court", profile.get("court")),
        ("Courthouse", profile.get("courthouse")),
        ("Presiding judge", profile.get("presidingJudge")),
        ("Hearing date", profile.get("hearingDate")),
        ("Defence counsel", profile.get("defenceCounsel")),
        ("Crown counsel", profile.get("crownCounsel")),
    ]

    for key, value in fields:
        if not value:
            continue
        _docx_add_kv_row(table, key, str(value))

    doc.add_page_break()


def _docx_section_documents(doc, payload):
    _docx_add_section_title(doc, "Documents reviewed")
    docs = payload.get("documents", []) or []

    if not docs:
        _docx_add_cited_paragraph(doc, "No documents have been uploaded for this matter.")
        doc.add_page_break()
        return

    for d in docs:
        p = doc.add_paragraph()
        run = p.add_run(d.get("filename", "Untitled document"))
        run.bold = True
        summary = d.get("summary") or d.get("analysis_summary")
        if summary:
            _docx_add_cited_paragraph(doc, summary)

    doc.add_page_break()


def _docx_section_intake(doc, payload):
    _docx_add_section_title(doc, "Intake interview - extracted state")
    intake = payload.get("intake_extracted", {}) or {}

    if not intake:
        _docx_add_cited_paragraph(
            doc,
            "No intake interview has been conducted, or no structured state was "
            "extracted from the conversation."
        )
        doc.add_page_break()
        return

    intake_fields = intake.get("fields") or intake
    if isinstance(intake_fields, dict) and intake_fields:
        table = doc.add_table(rows=0, cols=2)
        for key, value in intake_fields.items():
            if value in (None, "", []):
                continue
            display_key = key.replace("_", " ").capitalize()
            if isinstance(value, list):
                display_val = ", ".join(str(v) for v in value)
            else:
                display_val = str(value)
            _docx_add_kv_row(table, display_key, display_val)

    summary = intake.get("summary")
    if summary:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Interviewer's note")
        run.bold = True
        _docx_add_cited_paragraph(doc, summary)

    doc.add_page_break()


def _docx_section_risk(doc, payload):
    _docx_add_section_title(doc, "Risk state")
    evidence = payload.get("evidence", {}) or {}
    node_labels = payload.get("node_labels", {}) or {}

    if not evidence:
        _docx_add_cited_paragraph(
            doc,
            "No soft-evidence has been entered for this case. All nodes remain at "
            "their prior values."
        )
        doc.add_page_break()
        return

    _docx_add_cited_paragraph(
        doc,
        f"The following soft-evidence values are currently configured. "
        f"{len(evidence)} node(s) carry non-default values. Each value reflects the "
        f"degree of evidential support for the proposition that the named factor is present."
    )

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Node"
    hdr[1].text = "Label"
    hdr[2].text = "Value"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for node_id, value in sorted(evidence.items()):
        if value in (None, ""):
            continue
        row = table.add_row()
        row.cells[0].text = str(node_id)
        row.cells[1].text = node_labels.get(node_id, "-")
        if isinstance(value, float):
            row.cells[2].text = f"{value:.2f}"
        else:
            row.cells[2].text = str(value)

    doc.add_page_break()


def _docx_section_gladue(doc, payload):
    _docx_add_section_title(doc, "Gladue submission")
    gladue = payload.get("gladue", {}) or {}
    narrative = gladue.get("narrative") or ""
    factors = gladue.get("selected_factors", []) or []

    if not narrative and not factors:
        _docx_add_cited_paragraph(
            doc,
            "No Gladue submission has been generated for this matter. "
            "Where the offender has self-identified as Indigenous, sentencing courts are "
            "required to give particular attention to the circumstances of Indigenous "
            "offenders under s.718.2(e) and the Gladue/Ipeelee/Sharma line of authority."
        )
        doc.add_page_break()
        return

    if factors:
        p = doc.add_paragraph()
        run = p.add_run(f"Factors identified ({len(factors)})")
        run.bold = True
        for f in factors:
            label = f.get("label") or f.get("id") or "-"
            cat = f.get("category", "")
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(label)
            if cat:
                run = bullet.add_run(f"  ({cat})")
                run.italic = True
        doc.add_paragraph()

    if narrative:
        for para in narrative.split("\n\n"):
            para = para.strip()
            if para:
                _docx_add_cited_paragraph(doc, para)

    doc.add_page_break()


def _docx_section_sce(doc, payload):
    _docx_add_section_title(doc, "Social context evidence - Morris / Ellis")
    sce = payload.get("sce", {}) or {}
    narrative = sce.get("narrative") or ""
    factors = sce.get("selected_factors", []) or []

    if not narrative and not factors:
        _docx_add_cited_paragraph(
            doc,
            "No social context evidence submission has been generated for this matter. "
            "Where systemic and background factors of anti-Black racism are present, the "
            "framework in R v Morris, R v Anderson, and R v Ellis requires the sentencing "
            "judge to consider those factors as part of the proportionate sentence."
        )
        doc.add_page_break()
        return

    if factors:
        p = doc.add_paragraph()
        run = p.add_run(f"Factors identified ({len(factors)})")
        run.bold = True
        for f in factors:
            label = f.get("label") or f.get("id") or "-"
            cat = f.get("category", "")
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(label)
            if cat:
                run = bullet.add_run(f"  ({cat})")
                run.italic = True
        doc.add_paragraph()

    if narrative:
        for para in narrative.split("\n\n"):
            para = para.strip()
            if para:
                _docx_add_cited_paragraph(doc, para)

    doc.add_page_break()


def _docx_section_authorities(doc, payload):
    _docx_add_section_title(doc, "Cited doctrinal authorities")
    gladue_narr = (payload.get("gladue") or {}).get("narrative", "")
    sce_narr = (payload.get("sce") or {}).get("narrative", "")
    authorities = extract_authorities(gladue_narr, sce_narr)

    if not authorities:
        _docx_add_cited_paragraph(
            doc,
            "No doctrinal authorities have been cited in the generated submissions."
        )
        return

    for label in authorities:
        full = FULL_CITATIONS.get(label, label)
        bullet = doc.add_paragraph(style="List Bullet")
        run = bullet.add_run(full)
        run.italic = True


def build_audit_docx(payload, sections=None):
    """Render a DOCX audit report and return its bytes."""
    sections = sections or ALL_SECTIONS
    doc = Document()
    _docx_set_default_font(doc)

    for s in doc.sections:
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)

    section_fns = {
        "title": _docx_section_title_page,
        "executive_summary": _docx_section_executive_summary,
        "profile": _docx_section_profile,
        "documents": _docx_section_documents,
        "intake": _docx_section_intake,
        "risk": _docx_section_risk,
        "gladue": _docx_section_gladue,
        "sce": _docx_section_sce,
        "authorities": _docx_section_authorities,
    }
    for section_id in ALL_SECTIONS:
        if section_id in sections and section_id in section_fns:
            section_fns[section_id](doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================================
# PDF BUILDER (reportlab)
# ============================================================================

def _pdf_styles():
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "PARVIS_Title", parent=base["Title"],
            fontName="Times-Bold", fontSize=42, leading=46,
            alignment=TA_CENTER, spaceAfter=0,
        ),
        "Subtitle": ParagraphStyle(
            "PARVIS_Subtitle", parent=base["Normal"],
            fontName="Times-Italic", fontSize=13, leading=18,
            alignment=TA_CENTER, spaceAfter=24,
        ),
        "AuditTitle": ParagraphStyle(
            "PARVIS_AuditTitle", parent=base["Normal"],
            fontName="Times-Bold", fontSize=22, leading=28,
            alignment=TA_CENTER, spaceAfter=20,
        ),
        "Heading": ParagraphStyle(
            "PARVIS_Heading", parent=base["Heading1"],
            fontName="Times-Bold", fontSize=18, leading=22,
            alignment=TA_LEFT, spaceAfter=12, spaceBefore=0,
            textColor=colors.HexColor("#1a1a1a"),
        ),
        "Body": ParagraphStyle(
            "PARVIS_Body", parent=base["BodyText"],
            fontName="Times-Roman", fontSize=11, leading=15,
            alignment=TA_JUSTIFY, spaceAfter=8,
        ),
        "BodyCenter": ParagraphStyle(
            "PARVIS_BodyCenter", parent=base["BodyText"],
            fontName="Times-Roman", fontSize=12, leading=16,
            alignment=TA_CENTER, spaceAfter=6,
        ),
        "MetaCenter": ParagraphStyle(
            "PARVIS_MetaCenter", parent=base["BodyText"],
            fontName="Times-Italic", fontSize=9, leading=12,
            alignment=TA_CENTER, spaceAfter=0,
        ),
        "Bullet": ParagraphStyle(
            "PARVIS_Bullet", parent=base["BodyText"],
            fontName="Times-Roman", fontSize=11, leading=14,
            alignment=TA_LEFT, leftIndent=20, bulletIndent=8, spaceAfter=4,
        ),
        "Label": ParagraphStyle(
            "PARVIS_Label", parent=base["BodyText"],
            fontName="Times-Bold", fontSize=11, leading=14,
            alignment=TA_LEFT, spaceAfter=4, spaceBefore=8,
        ),
    }


def _escape_pdf(text):
    if text is None:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _render_cited_text_pdf(text):
    """Render [gladue:N] markers as bold italic inline reportlab markup."""
    parts = []
    for kind, value, para in parse_citations(text):
        if kind == "text":
            parts.append(_escape_pdf(value))
        else:
            label = _escape_pdf(value)
            if para:
                parts.append(f"<b><i>{label} para {_escape_pdf(para)}</i></b>")
            else:
                parts.append(f"<b><i>{label}</i></b>")
    return "".join(parts)


def _pdf_section_title_page(story, styles, payload):
    story.append(Spacer(1, 5*cm))
    story.append(Paragraph("PARVIS", styles["Title"]))
    story.append(Paragraph(
        "Probabilistic and Analytical Reasoning Virtual Intelligence System",
        styles["Subtitle"],
    ))
    story.append(Spacer(1, 2*cm))

    case_ref = payload.get("case_reference") or "Case reference not set"
    profile = payload.get("profile", {}) or {}
    given = profile.get("givenName", "") or ""
    family = profile.get("familyName", "") or ""
    full_name = f"{given} {family}".strip() or "Subject"
    jurisdiction = profile.get("jurisdiction", "") or ""
    court = profile.get("court", "") or ""
    charge = profile.get("primaryCharge", "") or ""
    generated = payload.get("generated_at") or datetime.utcnow().isoformat()

    story.append(Paragraph("Audit Report", styles["AuditTitle"]))

    forum_line = f"{court} - {jurisdiction}".strip(" -")
    for line in [case_ref, full_name, charge, forum_line]:
        if line:
            story.append(Paragraph(_escape_pdf(line), styles["BodyCenter"]))

    story.append(Spacer(1, 4*cm))
    story.append(Paragraph(
        f"Generated {generated[:10]} | {VERSION}",
        styles["MetaCenter"],
    ))
    story.append(PageBreak())


def _pdf_section_executive_summary(story, styles, payload):
    story.append(Paragraph("Executive summary", styles["Heading"]))
    inference = payload.get("inference", {}) or {}
    posterior = inference.get("do_posterior")

    if posterior is None:
        story.append(Paragraph(
            "Posterior probability of designation has not been computed for this case. "
            "Please refer to the Risk State section for current evidence values.",
            styles["Body"],
        ))
    else:
        pct = posterior * 100
        zone = (
            "above the beyond-reasonable-doubt threshold (&gt;=87%)"
            if pct >= 87
            else "below the beyond-reasonable-doubt threshold (&lt;87%)"
        )
        text = (
            f"Posterior probability that the predicate findings for designation "
            f"are established on the evidence currently configured: <b>{pct:.1f}%</b>. "
            f"This is {zone}. The standard of proof at a s.753 hearing is beyond "
            f"a reasonable doubt; PARVIS operationalises BRD at 87% on the basis "
            f"of empirical literature on judicial interpretation of that standard "
            f"(<i>R v Lifchus</i> and <i>R v Starr</i> describe BRD qualitatively, "
            f"without specifying a probability). PARVIS does not predict; it "
            f"surfaces the structure of the inference. This report is intended as "
            f"a doctrinally literate companion to the substantive submissions of counsel."
        )
        story.append(Paragraph(text, styles["Body"]))

    drivers_up = inference.get("drivers_up", []) or []
    drivers_down = inference.get("drivers_down", []) or []

    if drivers_up or drivers_down:
        story.append(Paragraph("Drivers", styles["Label"]))
        if drivers_up:
            line = "Increasing posterior: " + ", ".join(
                f"{_escape_pdf(d.get('label', d.get('nodeId', '?')))} ({d.get('delta', 0):+.1%})"
                for d in drivers_up[:5]
            )
            story.append(Paragraph(line, styles["Body"]))
        if drivers_down:
            line = "Decreasing posterior: " + ", ".join(
                f"{_escape_pdf(d.get('label', d.get('nodeId', '?')))} ({d.get('delta', 0):+.1%})"
                for d in drivers_down[:5]
            )
            story.append(Paragraph(line, styles["Body"]))

    story.append(PageBreak())


def _kv_table(rows):
    """Helper: build a two-column key/value Table flowable."""
    table = Table(rows, colWidths=[5*cm, 11*cm])
    table.setStyle(TableStyle([
        ("FONT", (0, 0), (0, -1), "Times-Bold", 10),
        ("FONT", (1, 0), (1, -1), "Times-Roman", 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
    ]))
    return table


def _pdf_section_profile(story, styles, payload):
    story.append(Paragraph("Case profile", styles["Heading"]))
    profile = payload.get("profile", {}) or {}

    if not profile:
        story.append(Paragraph("No case profile data captured.", styles["Body"]))
        story.append(PageBreak())
        return

    fields = [
        ("Case reference", profile.get("caseReference")),
        ("Given name(s)", profile.get("givenName")),
        ("Family name", profile.get("familyName")),
        ("Date of birth", profile.get("dateOfBirth")),
        ("Pronouns", profile.get("pronouns")),
        ("Indigenous identity", INDIGENOUS_LABELS.get(profile.get("indigenousIdentity", "none"), "-")),
        ("Nation / community", profile.get("nationCommunity")),
        ("Place of origin", profile.get("placeOfOrigin")),
        ("Current residence", profile.get("currentResidence")),
        ("FASD diagnosis", profile.get("fasdDiagnosis")),
        ("Cognitive assessment", profile.get("cognitiveAssess")),
        ("Mental health diagnosis", profile.get("mentalHealthDx")),
        ("Primary charge", profile.get("primaryCharge")),
        ("Additional charges", profile.get("additionalCharges")),
        ("s.753 application", profile.get("s753Application")),
        ("Crown position", profile.get("crownPosition")),
        ("Defence position", profile.get("defencePosition")),
        ("Jurisdiction", profile.get("jurisdiction")),
        ("Court", profile.get("court")),
        ("Courthouse", profile.get("courthouse")),
        ("Presiding judge", profile.get("presidingJudge")),
        ("Hearing date", profile.get("hearingDate")),
        ("Defence counsel", profile.get("defenceCounsel")),
        ("Crown counsel", profile.get("crownCounsel")),
    ]
    rows = [[_escape_pdf(k), _escape_pdf(v)] for k, v in fields if v]
    if rows:
        story.append(_kv_table(rows))
    story.append(PageBreak())


def _pdf_section_documents(story, styles, payload):
    story.append(Paragraph("Documents reviewed", styles["Heading"]))
    docs = payload.get("documents", []) or []

    if not docs:
        story.append(Paragraph("No documents have been uploaded for this matter.", styles["Body"]))
        story.append(PageBreak())
        return

    for d in docs:
        story.append(Paragraph(_escape_pdf(d.get("filename", "Untitled document")), styles["Label"]))
        summary = d.get("summary") or d.get("analysis_summary")
        if summary:
            story.append(Paragraph(_render_cited_text_pdf(summary), styles["Body"]))

    story.append(PageBreak())


def _pdf_section_intake(story, styles, payload):
    story.append(Paragraph("Intake interview - extracted state", styles["Heading"]))
    intake = payload.get("intake_extracted", {}) or {}

    if not intake:
        story.append(Paragraph(
            "No intake interview has been conducted, or no structured state was "
            "extracted from the conversation.",
            styles["Body"],
        ))
        story.append(PageBreak())
        return

    intake_fields = intake.get("fields") or intake
    if isinstance(intake_fields, dict) and intake_fields:
        rows = []
        for key, value in intake_fields.items():
            if value in (None, "", []):
                continue
            display_key = key.replace("_", " ").capitalize()
            if isinstance(value, list):
                display_val = ", ".join(str(v) for v in value)
            else:
                display_val = str(value)
            rows.append([_escape_pdf(display_key), _escape_pdf(display_val)])
        if rows:
            story.append(_kv_table(rows))

    summary = intake.get("summary")
    if summary:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("Interviewer's note", styles["Label"]))
        story.append(Paragraph(_render_cited_text_pdf(summary), styles["Body"]))

    story.append(PageBreak())


def _pdf_section_risk(story, styles, payload):
    story.append(Paragraph("Risk state", styles["Heading"]))
    evidence = payload.get("evidence", {}) or {}
    node_labels = payload.get("node_labels", {}) or {}

    if not evidence:
        story.append(Paragraph(
            "No soft-evidence has been entered for this case. All nodes remain at "
            "their prior values.",
            styles["Body"],
        ))
        story.append(PageBreak())
        return

    story.append(Paragraph(
        f"The following soft-evidence values are currently configured. "
        f"{len(evidence)} node(s) carry non-default values. Each value reflects the "
        f"degree of evidential support for the proposition that the named factor is present.",
        styles["Body"],
    ))

    rows = [["Node", "Label", "Value"]]
    for node_id, value in sorted(evidence.items()):
        if value in (None, ""):
            continue
        display_val = f"{value:.2f}" if isinstance(value, float) else str(value)
        rows.append([
            _escape_pdf(node_id),
            _escape_pdf(node_labels.get(node_id, "-")),
            _escape_pdf(display_val),
        ])

    if len(rows) > 1:
        table = Table(rows, colWidths=[3*cm, 10*cm, 3*cm])
        table.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, 0), "Times-Bold", 10),
            ("FONT", (0, 1), (-1, -1), "Times-Roman", 10),
            ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#1a1a1a")),
            ("LINEBELOW", (0, 1), (-1, -1), 0.25, colors.HexColor("#dddddd")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)

    story.append(PageBreak())


def _pdf_section_gladue(story, styles, payload):
    story.append(Paragraph("Gladue submission", styles["Heading"]))
    gladue = payload.get("gladue", {}) or {}
    narrative = gladue.get("narrative") or ""
    factors = gladue.get("selected_factors", []) or []

    if not narrative and not factors:
        story.append(Paragraph(
            "No Gladue submission has been generated for this matter. "
            "Where the offender has self-identified as Indigenous, sentencing courts are "
            "required to give particular attention to the circumstances of Indigenous "
            "offenders under s.718.2(e) and the Gladue/Ipeelee/Sharma line of authority.",
            styles["Body"],
        ))
        story.append(PageBreak())
        return

    if factors:
        story.append(Paragraph(f"Factors identified ({len(factors)})", styles["Label"]))
        for f in factors:
            label = _escape_pdf(f.get("label") or f.get("id") or "-")
            cat = _escape_pdf(f.get("category", ""))
            text = f"&bull; {label}"
            if cat:
                text += f"  <i>({cat})</i>"
            story.append(Paragraph(text, styles["Bullet"]))
        story.append(Spacer(1, 0.4*cm))

    if narrative:
        for para in narrative.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(_render_cited_text_pdf(para), styles["Body"]))

    story.append(PageBreak())


def _pdf_section_sce(story, styles, payload):
    story.append(Paragraph("Social context evidence - Morris / Ellis", styles["Heading"]))
    sce = payload.get("sce", {}) or {}
    narrative = sce.get("narrative") or ""
    factors = sce.get("selected_factors", []) or []

    if not narrative and not factors:
        story.append(Paragraph(
            "No social context evidence submission has been generated for this matter. "
            "Where systemic and background factors of anti-Black racism are present, the "
            "framework in R v Morris, R v Anderson, and R v Ellis requires the sentencing "
            "judge to consider those factors as part of the proportionate sentence.",
            styles["Body"],
        ))
        story.append(PageBreak())
        return

    if factors:
        story.append(Paragraph(f"Factors identified ({len(factors)})", styles["Label"]))
        for f in factors:
            label = _escape_pdf(f.get("label") or f.get("id") or "-")
            cat = _escape_pdf(f.get("category", ""))
            text = f"&bull; {label}"
            if cat:
                text += f"  <i>({cat})</i>"
            story.append(Paragraph(text, styles["Bullet"]))
        story.append(Spacer(1, 0.4*cm))

    if narrative:
        for para in narrative.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(_render_cited_text_pdf(para), styles["Body"]))

    story.append(PageBreak())


def _pdf_section_authorities(story, styles, payload):
    story.append(Paragraph("Cited doctrinal authorities", styles["Heading"]))
    gladue_narr = (payload.get("gladue") or {}).get("narrative", "")
    sce_narr = (payload.get("sce") or {}).get("narrative", "")
    authorities = extract_authorities(gladue_narr, sce_narr)

    if not authorities:
        story.append(Paragraph(
            "No doctrinal authorities have been cited in the generated submissions.",
            styles["Body"],
        ))
        return

    for label in authorities:
        full = FULL_CITATIONS.get(label, label)
        story.append(Paragraph(f"&bull; <i>{_escape_pdf(full)}</i>", styles["Bullet"]))


def _pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Italic", 8)
    canvas.setFillColor(colors.HexColor("#888888"))
    page_num = canvas.getPageNumber()
    canvas.drawCentredString(
        A4[0] / 2, 1.2*cm,
        f"PARVIS Audit Report | page {page_num}"
    )
    canvas.restoreState()


def build_audit_pdf(payload, sections=None):
    """Render a PDF audit report and return its bytes."""
    sections = sections or ALL_SECTIONS
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    styles = _pdf_styles()
    story = []

    section_fns = {
        "title": _pdf_section_title_page,
        "executive_summary": _pdf_section_executive_summary,
        "profile": _pdf_section_profile,
        "documents": _pdf_section_documents,
        "intake": _pdf_section_intake,
        "risk": _pdf_section_risk,
        "gladue": _pdf_section_gladue,
        "sce": _pdf_section_sce,
        "authorities": _pdf_section_authorities,
    }
    for section_id in ALL_SECTIONS:
        if section_id in sections and section_id in section_fns:
            section_fns[section_id](story, styles, payload)

    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return buf.getvalue()
